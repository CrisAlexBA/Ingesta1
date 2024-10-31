import os
import boto3
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# Configuración
BUCKET_NAME = 'uq-datalake-test'  # Nombre del bucket
RUTA_PREFIJO = 'archivos/'  # Carpeta donde se guardarán los archivos

def crear_pdf(nombre_archivo, contenido):
    """Crea un archivo PDF con el contenido especificado."""
    c = canvas.Canvas(nombre_archivo, pagesize=letter)
    c.drawString(100, 750, contenido)
    c.save()

def crear_docx(nombre_archivo, contenido):
    """Crea un archivo DOCX con el contenido especificado."""
    doc = Document()
    doc.add_heading('Documento de Prueba', level=1)
    doc.add_paragraph(contenido)
    doc.save(nombre_archivo)

def subir_a_s3(nombre_archivo, ruta_s3):
    """Sube un archivo a S3."""
    s3 = boto3.client('s3')
    s3.upload_file(nombre_archivo, BUCKET_NAME, ruta_s3)

def generar_archivos():
    """Genera y sube archivos PDF y DOCX de prueba a S3."""
    for i in range(1, 4):  # Generar 3 archivos de cada tipo
        contenido = f"Este es el contenido del archivo de prueba número {i}."
        
        # Crear nombres de archivo
        nombre_pdf = f"archivo_prueba_{i}.pdf"
        nombre_docx = f"archivo_prueba_{i}.docx"
        
        # Crear los documentos
        crear_pdf(nombre_pdf, contenido)
        crear_docx(nombre_docx, contenido)
        
        # Subir los archivos a S3
        subir_a_s3(nombre_pdf, f"{RUTA_PREFIJO}{nombre_pdf}")
        subir_a_s3(nombre_docx, f"{RUTA_PREFIJO}{nombre_docx}")
        
        # Eliminar los archivos locales después de subirlos
        os.remove(nombre_pdf)
        os.remove(nombre_docx)

if __name__ == "__main__":
    generar_archivos()