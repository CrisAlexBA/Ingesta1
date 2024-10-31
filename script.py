import os
import sys
import random
import boto3
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# Agrega la ruta del directorio padre para acceder a 'utils'
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Importar funciones y módulos de utils
from s3_utils.manejador_metadata import generar_metadata
from utils.funciones_auxiliares import obtener_semestre  # Asegúrate de que esta función esté definida

def generar_grupo_aleatorio():
    """Selecciona un grupo aleatorio de las opciones predefinidas."""
    grupos = ['ceifi', 'externo', 'grid', 'sinfoci']
    return random.choice(grupos)

# Configuración
BUCKET_NAME = 'uq-datalake-test'
RUTA_PREFIJO = 'archivos/'

def crear_pdf(nombre_archivo, contenido):
    c = canvas.Canvas(nombre_archivo, pagesize=letter)
    c.drawString(100, 750, contenido)
    c.save()

def crear_docx(nombre_archivo, contenido):
    doc = Document()
    doc.add_heading('Documento de Prueba', level=1)
    doc.add_paragraph(contenido)
    doc.save(nombre_archivo)

def subir_a_s3(nombre_archivo, ruta_s3, metadata):
    s3 = boto3.client('s3')
    s3.upload_file(nombre_archivo, BUCKET_NAME, ruta_s3, ExtraArgs={"Metadata": metadata})

def generar_archivos():
    for i in range(1, 4):  # Generar 3 archivos de cada tipo
        contenido = f"Este es el contenido del archivo de prueba número {i}."
        nombre_pdf = f"archivo_prueba_{i}.pdf"
        nombre_docx = f"archivo_prueba_{i}.docx"
        
        semestre = obtener_semestre()  # Asegúrate de que esta función esté definida
        grupo = generar_grupo_aleatorio()  # Obtener grupo aleatorio
        metadata = generar_metadata(semestre, grupo)

        crear_pdf(nombre_pdf, contenido)
        crear_docx(nombre_docx, contenido)
        
        # Subir archivos a S3 en la carpeta correspondiente al grupo
        subir_a_s3(nombre_pdf, f"{RUTA_PREFIJO}{nombre_pdf}", metadata)
        subir_a_s3(nombre_docx, f"{RUTA_PREFIJO}{nombre_docx}", metadata)
        
        # Eliminar los archivos locales después de subirlos
        os.remove(nombre_pdf)
        os.remove(nombre_docx)

if __name__ == "__main__":
    generar_archivos()
